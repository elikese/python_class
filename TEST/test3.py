import os
import time
import re
import requests
from bs4 import BeautifulSoup
from fake_useragent import UserAgent
from docx import Document

# ===============================
# 🧭 설정 영역
# ===============================
BASE_URL = "https://overseas.mofa.go.kr/de-ko/brd/m_7204/view.do?seq={seq}&page={page}"

START_PAGE = 1  # 시작 페이지
END_PAGE = 2  # 마지막 페이지
START_SEQ = 2975087  # 1페이지 첫 게시글 seq 값
POSTS_PER_PAGE = 10  # 페이지당 게시글 수 (관찰 결과 약 10개)

SAVE_DIR = "crawled_posts"  # 저장 폴더명
os.makedirs(SAVE_DIR, exist_ok=True)


# ===============================
# 🧰 함수 정의
# ===============================
def fetch_html(url):
    """HTML을 요청하고 BeautifulSoup 객체로 반환"""
    headers = {"User-Agent": UserAgent().random}
    res = requests.get(url, headers=headers, timeout=10)
    res.raise_for_status()
    return BeautifulSoup(res.text, "lxml")


def parse_post(soup):
    """게시글 HTML에서 제목, 작성자, 작성일, 본문 추출"""
    title = soup.select_one("div.board_detail > div.bo_head > h2").get_text(strip=True)
    author = soup.select_one("div.board_detail > div.bo_head > dl > dd:nth-of-type(1)").get_text(strip=True)
    date = soup.select_one("div.board_detail > div.bo_head > dl > dd:nth-of-type(2)").get_text(strip=True)
    content_html = soup.select_one("div.board_detail > div.bo_con")
    content_text = content_html.get_text("\n", strip=True)
    return title, author, date, content_text


def save_to_docx(title, author, date, content):
    """게시글을 DOCX 파일로 저장"""
    safe_title = re.sub(r'[\\/:*?"<>|]', "_", title)  # 파일명에서 불가능한 문자 제거
    filename = f"{safe_title}_{date}.docx"
    filepath = os.path.join(SAVE_DIR, filename)

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"작성자: {author}")
    doc.add_paragraph(f"작성일: {date}")
    doc.add_paragraph("")
    doc.add_paragraph(content)
    doc.save(filepath)
    print(f"💾 저장 완료: {filename}")


# ===============================
# 🚀 크롤링 실행
# ===============================
if __name__ == "__main__":
    current_seq = START_SEQ

    for page in range(START_PAGE, END_PAGE + 1):
        print(f"\n📄 페이지 {page} 처리 중...")

        for i in range(POSTS_PER_PAGE):
            seq = current_seq - i - (page - START_PAGE) * POSTS_PER_PAGE
            url = BASE_URL.format(seq=seq, page=page)
            print(f"→ 요청 중: {url}")

            try:
                soup = fetch_html(url)
                title, author, date, content = parse_post(soup)
                save_to_docx(title, author, date, content)
            except Exception as e:
                print(f"❌ 오류 발생 ({url}): {e}")

            time.sleep(1.5)  # 서버 과부하 방지용 딜레이
