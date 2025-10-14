import requests
from bs4 import BeautifulSoup
from docx import Document
import os
import re
import time


# ==========================================
# ⚙️ 설정
# ==========================================

START_PAGE = 1  # 시작 페이지
END_PAGE = 2  # 마지막 페이지
START_SEQ = 2975087  # 1페이지 첫 게시글 seq
POSTS_PER_PAGE = 10  # 페이지당 게시글 수
SAVE_DIR = "output"  # 결과 저장 폴더

BASE_URL = "https://overseas.mofa.go.kr/de-ko/brd/m_7204/view.do?seq={seq}&page={page}"

os.makedirs(SAVE_DIR, exist_ok=True)


# ==========================================
# 🧩 유틸리티 함수
# ==========================================


def clean_text(text: str) -> str:
    """불필요한 공백 제거"""
    return re.sub(r"\s+", " ", text).strip()


def safe_filename(text: str) -> str:
    """파일명으로 부적합한 문자 제거"""
    return re.sub(r'[\\/*?:"<>|]', "", text)


# ==========================================
# 🧠 게시글 상세 데이터 추출
# ==========================================


def parse_post(seq: int, page: int):
    """게시글 상세 페이지에서 제목, 작성자, 작성일, 본문 추출"""
    url = BASE_URL.format(seq=seq, page=page)
    res = requests.get(url)
    if res.status_code != 200:
        print(f"❌ 요청 실패: {url}")
        return None

    soup = BeautifulSoup(res.text, "lxml")

    # 제목
    title_tag = soup.select_one(".bo_head h2")
    title = clean_text(title_tag.get_text()) if title_tag else f"게시글_{seq}"

    # 작성자 / 작성일
    info_tags = soup.select(".bo_head dl dd")
    author = clean_text(info_tags[0].get_text()) if len(info_tags) > 0 else "작성자 없음"
    date = clean_text(info_tags[1].get_text()) if len(info_tags) > 1 else "날짜 없음"

    # 본문
    body_div = soup.find("div", class_="bo_con")
    if body_div:
        body_text = body_div.get_text(separator="\n", strip=True)
        body_text = re.sub(r"\s+", " ", body_text).strip()
    else:
        body_text = "본문 없음"

    return {
        "url": url,
        "title": title,
        "author": author,
        "date": date,
        "content": body_text,
    }


# ==========================================
# 💾 DOCX 파일 저장
# ==========================================


def save_to_docx(post_data: dict):
    """게시글을 DOCX 파일로 저장"""
    document = Document()
    document.add_heading(post_data["title"], level=1)
    document.add_paragraph(f"URL: {post_data['url']}")
    document.add_paragraph(f"작성자: {post_data['author']}")
    document.add_paragraph(f"작성일: {post_data['date']}")
    document.add_paragraph("")  # 공백 줄
    document.add_paragraph(post_data["content"])

    safe_title = safe_filename(post_data["title"])
    safe_date = safe_filename(post_data["date"])
    filename = f"{safe_title}_{safe_date}.docx"
    path = os.path.join(SAVE_DIR, filename)

    document.save(path)
    print(f"✅ 저장 완료: {path}")


# ==========================================
# 🚀 메인 실행 로직
# ==========================================


def main():
    print(f"📄 {START_PAGE}페이지부터 {END_PAGE}페이지까지 크롤링 시작...\n")

    for page in range(START_PAGE, END_PAGE + 1):
        print(f"🧭 {page}페이지 처리 중...")

        start_seq = START_SEQ - (page - 1) * POSTS_PER_PAGE
        for i in range(POSTS_PER_PAGE):
            seq = start_seq - i
            print(f"  ▶ 게시글 seq={seq} 요청 중...")

            try:
                post = parse_post(seq, page)
                if post:
                    save_to_docx(post)
                time.sleep(0.5)  # 서버 부하 방지
            except Exception as e:
                print(f"❌ 오류 발생 (seq={seq}): {e}")

    print("\n🎉 모든 페이지 크롤링 완료!")


if __name__ == "__main__":
    main()
