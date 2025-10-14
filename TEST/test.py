import requests
from bs4 import BeautifulSoup
from docx import Document
import os
import re
import time


# ==========================================
# ⚙️ 설정 영역
# ==========================================

START_PAGE = 1  # 시작 페이지
END_PAGE = 2  # 마지막 페이지
SAVE_DIR = "output"  # 결과 저장 폴더

BASE_URL = "https://overseas.mofa.go.kr"
BOARD_PATH = "/de-ko/brd/m_7204"  # 독일 대사관 게시판 경로

LIST_URL = f"{BASE_URL}{BOARD_PATH}/list.do?page={{page}}"
VIEW_URL = f"{BASE_URL}{BOARD_PATH}/view.do?seq={{seq}}&page={{page}}"

os.makedirs(SAVE_DIR, exist_ok=True)


# ==========================================
# 🧩 유틸리티 함수
# ==========================================


def clean_text(text: str) -> str:
    """HTML에서 불필요한 공백 제거"""
    return re.sub(r"\s+", " ", text).strip()


def safe_filename(text: str) -> str:
    """파일명으로 부적합한 문자 제거"""
    return re.sub(r'[\\/*?:"<>|]', "", text)


# ==========================================
# 📋 게시글 목록 추출 함수
# ==========================================


def get_post_links(page: int):
    """지정된 페이지에서 게시글 seq 목록 추출"""
    url = LIST_URL.format(page=page)
    res = requests.get(url)
    res.raise_for_status()
    soup = BeautifulSoup(res.text, "lxml")

    post_links = []
    # 게시판 목록의 제목 영역에서 게시글 링크 추출
    for a_tag in soup.select("table.board_list td.title a"):
        href = a_tag.get("href", "")
        m = re.search(r"seq=(\d+)", href)
        if m:
            seq = m.group(1)
            post_links.append(seq)

    # 중복 제거
    post_links = list(dict.fromkeys(post_links))
    return post_links


# ==========================================
# 🧠 게시글 상세 데이터 추출 함수
# ==========================================


def parse_post(seq: str, page: int):
    """게시글 상세 페이지에서 데이터 추출"""
    url = VIEW_URL.format(seq=seq, page=page)
    res = requests.get(url)
    res.raise_for_status()
    soup = BeautifulSoup(res.text, "lxml")

    # 제목
    title_tag = soup.select_one(".bo_head h2")
    title = clean_text(title_tag.get_text()) if title_tag else f"게시글_{seq}"

    # 작성자 / 작성일
    info_tags = soup.select(".bo_head dl dd")
    author = clean_text(info_tags[0].get_text()) if len(info_tags) > 0 else "작성자 없음"
    date = clean_text(info_tags[1].get_text()) if len(info_tags) > 1 else "날짜 없음"

    # 본문
    content_tag = soup.select_one(".bo_con .se-contents")
    if content_tag:
        # <p> 단위로 줄바꿈 처리
        paragraphs = [clean_text(p.get_text()) for p in content_tag.find_all("p") if clean_text(p.get_text())]
        content = "\n".join(paragraphs)
    else:
        content = "본문 없음"

    return {
        "title": title,
        "author": author,
        "date": date,
        "content": content,
    }


# ==========================================
# 💾 DOCX 파일 저장 함수
# ==========================================


def save_to_docx(post_data: dict):
    """게시글 데이터를 .docx 파일로 저장"""
    document = Document()
    document.add_heading(post_data["title"], level=1)
    document.add_paragraph(f"작성자: {post_data['author']}")
    document.add_paragraph(f"작성일: {post_data['date']}")
    document.add_paragraph("")  # 공백 줄
    document.add_paragraph(post_data["content"])

    safe_title = safe_filename(post_data["title"])
    safe_date = safe_filename(post_data["date"])
    filename = f"{safe_title}_{safe_date}.docx"
    path = os.path.join(SAVE_DIR, filename)

    document.save(path)
    print(f"✅ 저장 완료: {path}")


# ==========================================
# 🚀 메인 실행 로직
# ==========================================


def main():
    print(f"📄 {START_PAGE}페이지부터 {END_PAGE}페이지까지 크롤링 시작...\n")

    for page in range(START_PAGE, END_PAGE + 1):
        print(f"🧭 {page}페이지 처리 중...")

        seq_list = get_post_links(page)

        if not seq_list:
            print(f"⚠️ {page}페이지에서 게시글을 찾지 못했습니다.")
            continue

        for seq in seq_list:
            try:
                post = parse_post(seq, page)
                save_to_docx(post)
                time.sleep(1)  # 서버 부하 방지
            except Exception as e:
                print(f"❌ 오류 발생 (seq={seq}): {e}")

    print("\n🎉 모든 페이지 크롤링 완료!")


if __name__ == "__main__":
    main()
